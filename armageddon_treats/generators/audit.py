"""
Generate professionally formatted Word document for the Armageddon Treats
Business Plan Audit (main report).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ── Brand palette ──────────────────────────────────────────────
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY   = RGBColor(0x2D, 0x2D, 0x2D)
MID_GRAY    = RGBColor(0x4A, 0x4A, 0x4A)
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)
ACCENT      = RGBColor(0xC0, 0x39, 0x2B)   # deep red / apocalyptic accent
ACCENT_DARK = RGBColor(0x96, 0x2D, 0x22)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "1A1A1A"
TABLE_ALT_BG    = "F7F7F7"
TABLE_BORDER    = "CCCCCC"

FONT_BODY    = "Calibri"
FONT_HEADING = "Calibri"


# ── Helpers ────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", sz="4"):
    """Set thin borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def style_header_row(row, bg=TABLE_HEADER_BG, font_color=WHITE):
    """Style a table header row."""
    for cell in row.cells:
        set_cell_shading(cell, bg)
        set_cell_borders(cell, "999999")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = font_color
                run.font.size = Pt(9.5)
                run.font.name = FONT_HEADING


def style_data_row(row, idx):
    """Style an alternating data row."""
    bg = TABLE_ALT_BG if idx % 2 == 0 else "FFFFFF"
    for cell in row.cells:
        set_cell_shading(cell, bg)
        set_cell_borders(cell, TABLE_BORDER)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = FONT_BODY
                run.font.color.rgb = DARK_GRAY


def add_styled_table(doc, headers, rows, col_widths=None):
    """Insert a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
    style_header_row(hdr)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
        style_data_row(row, r_idx)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")  # spacer
    return table


def add_heading_styled(doc, text, level=1):
    """Add a heading with brand styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_HEADING
        run.font.color.rgb = BLACK if level == 1 else DARK_GRAY
    return h


def add_body_para(doc, text, bold=False, italic=False, space_after=Pt(8)):
    """Add a body paragraph with consistent styling."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GRAY
    run.bold = bold
    run.italic = italic
    return p


def setup_document():
    """Create a new document with standard margins, fonts, and header/footer."""
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(10)
    font.color.rgb = MID_GRAY

    # Heading styles
    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = FONT_HEADING
        hs.font.color.rgb = BLACK if i == 1 else DARK_GRAY
        hs.paragraph_format.space_before = Pt(18 if i == 1 else 14)
        hs.paragraph_format.space_after = Pt(6)
        if i == 1:
            hs.font.size = Pt(18)
        elif i == 2:
            hs.font.size = Pt(14)
        else:
            hs.font.size = Pt(12)

    return doc


def add_header_footer(doc, title_short):
    """Add header and footer to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(title_short)
        run.font.name = FONT_HEADING
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.small_caps = True

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("ARMAGEDDON TREATS  |  Veteran-Owned  |  Confidential")
        run.font.name = FONT_HEADING
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        run.font.small_caps = True


def add_cover_page(doc):
    """Create a branded title page."""
    # Spacer
    for _ in range(5):
        doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARMAGEDDON TREATS")
    run.font.name = FONT_HEADING
    run.font.size = Pt(36)
    run.font.color.rgb = BLACK
    run.bold = True
    run.font.small_caps = True

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    run2 = p2.add_run("Comprehensive Business Audit & Strategic Framework")
    run2.font.name = FONT_HEADING
    run2.font.size = Pt(16)
    run2.font.color.rgb = MID_GRAY

    # Divider line
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    run3 = p3.add_run("_" * 60)
    run3.font.color.rgb = ACCENT
    run3.font.size = Pt(10)

    # Meta info
    for _ in range(2):
        doc.add_paragraph("")

    meta_lines = [
        "Prepared for: Veteran Proprietor",
        "Classification: Confidential Business Document",
        "Fiscal Year: 2026",
        "Financing Vehicle: SBA 7(a) Microloan — $14,000",
        "Target Market: Santa Monica Farmers Markets, CA",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = MID_GRAY

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
#  MAIN DOCUMENT
# ══════════════════════════════════════════════════════════════
def build_main_document():
    doc = setup_document()
    add_cover_page(doc)
    add_header_footer(doc, "Armageddon Treats — Business Audit 2026")

    # ── SECTION 1: EXECUTIVE OVERVIEW ─────────────────────────
    add_heading_styled(doc, "1. Executive Overview", level=1)

    add_body_para(doc,
        "The convergence of experiential retail, shifting post-pandemic municipal regulations, "
        "and localized environmental ordinances creates a highly specific operational landscape "
        "for mobile food vendors in Southern California. This comprehensive audit examines the "
        "structural, financial, regulatory, and branding viability of \"Armageddon Treats,\" a "
        "proposed veteran-owned mobile dessert operation. The enterprise seeks to establish a "
        "highly visible presence at the Santa Monica Farmers Markets, capitalizing on a "
        "\"doomsday\" or apocalyptic theme while serving premium organic cotton candy and shaved "
        "ice. The venture is structured to be capitalized by a proposed $14,000 Small Business "
        "Administration (SBA) 7(a) microloan."
    )

    add_body_para(doc,
        "While the unit economics of the core products are exceptionally strong\u2014routinely "
        "yielding gross profit margins exceeding 90%\u2014the micro-capitalization of $14,000 "
        "presents severe operational constraints. Navigating Los Angeles County\u2019s strict "
        "Mobile Food Facility (MFF) health codes, the mandatory utilization of commercial "
        "commissary kitchens, and the highly competitive, narrowly opened Request for Proposal "
        "(RFP) processes governing Santa Monica\u2019s public markets requires flawless strategic "
        "execution. Furthermore, the business must seamlessly integrate its seemingly abrasive "
        "survivalist branding into a premium, family-oriented, and sustainability-focused "
        "municipal environment."
    )

    add_body_para(doc,
        "This report provides an exhaustive, multi-disciplinary evaluation of the market "
        "positioning, regulatory compliance pathways, supply chain logistics, and capital "
        "deployment strategies necessary to successfully execute the Armageddon Treats business "
        "plan in the 2026 fiscal environment."
    )

    # ── SECTION 2: MACRO-ENVIRONMENTAL ANALYSIS ───────────────
    add_heading_styled(doc, "2. Macro-Environmental & Consumer Branding Analysis", level=1)

    add_heading_styled(doc, "2.1 Foot Traffic and Consumer Dynamics in Santa Monica", level=2)

    add_body_para(doc,
        "Santa Monica represents one of the most premium retail environments in the United "
        "States, characterized by high-density foot traffic, affluent residential demographics, "
        "and a municipal government that aggressively promotes sustainability and organic "
        "agriculture. Post-pandemic recovery data indicates a robust and sustained return to "
        "physical retail, particularly in open-air markets and coastal commercial zones. By early "
        "2026, central commercial arteries such as the Third Street Promenade recorded foot "
        "traffic exceeding 329,343 monthly visitors, representing a stabilization and "
        "revitalization of the downtown districts."
    )

    add_body_para(doc,
        "The Santa Monica Farmers Markets operate across four distinct weekly iterations, "
        "fundamentally serving as community hubs that blend agricultural commerce with civic "
        "engagement. The Saturday Pico Boulevard market (located at Virginia Avenue Park) and "
        "the Sunday Main Street market (situated in the Heritage Square parking lot) serve as "
        "the primary strategic targets for prepared food vendors. The Sunday Main Street Market, "
        "operating from 8:30 AM to 1:30 PM, is a critical nexus of certified agricultural "
        "growers, specialty food vendors, and dense coastal foot traffic. It draws thousands of "
        "weekly visitors and features live music, complementary bike valets, and community "
        "activities, making it an ideal proving ground for a highly thematic, visually distinct "
        "food concept that relies on impulse purchases and pedestrian visibility."
    )

    # Target market summary table
    add_styled_table(doc,
        ["Market Location", "Day / Hours", "Strategic Value"],
        [
            ["Pico Boulevard (Virginia Ave Park)", "Saturday", "Secondary target; community-focused foot traffic"],
            ["Main Street (Heritage Square)", "Sunday / 8:30 AM \u2013 1:30 PM", "Primary target; highest density, live entertainment, coastal visitors"],
        ],
        col_widths=[2.2, 1.8, 2.8]
    )

    add_heading_styled(doc, "2.2 The Paradoxical Appeal of \u201cApocalyptic\u201d Branding in Luxury Markets", level=2)

    add_body_para(doc,
        "The foundational branding premise of Armageddon Treats relies on a \u201cdoomsday\u201d or "
        "apocalyptic theme, juxtaposed with the sale of brightly colored, nostalgic, and organic "
        "desserts. At first glance, survivalist aesthetics appear contradictory to the bright, "
        "wellness-oriented, and progressive ethos of Southern California\u2019s coastal markets. "
        "However, contemporary consumer behavior studies indicate a rising, mainstream "
        "receptiveness to apocalyptic aesthetics within premium and luxury markets."
    )

    add_body_para(doc,
        "The phenomenon of \u201cprepping\u201d\u2014traditionally associated with rural survivalism "
        "and off-grid living\u2014has transcended its fringe origins to become a mainstream "
        "cultural motif. Academic analyses of prepper culture reveal that the anticipation of "
        "systemic collapse allows consumers to mythically prefigure \u201chit-the-fan\u201d "
        "scenarios as a mechanism to ameliorate modern anxieties. In the context of the global "
        "crises experienced throughout the 2020s, an estimated 10% of all American households "
        "now actively engage in some form of disaster preparation, drawing interest from diverse "
        "demographics, including urban professionals and left-leaning communities who view "
        "self-sufficiency as a form of social resilience."
    )

    add_body_para(doc,
        "In the realm of corporate marketing and food branding, the adoption of a \u201cdark "
        "aesthetic\u201d\u2014characterized by raw industrial materials, dark color palettes, and "
        "utilitarian typography\u2014has been successfully utilized to signify authenticity and "
        "differentiate products from overly polished, traditional corporate branding. This "
        "approach has been prominently championed by the New Nordic food movement and high-end "
        "culinary institutions, which utilize dark visual palettes to communicate a return to "
        "foundational, unadulterated food experiences."
    )

    add_body_para(doc,
        "For Armageddon Treats, adopting a rugged, survivalist visual identity\u2014such as a "
        "matte-black trailer outfitted with industrial hardware and hazard-themed typography\u2014"
        "while serving brightly colored, USDA Organic spun sugar and shaved ice creates a "
        "highly shareable, visually arresting dissonance. This tension between a dystopian "
        "external aesthetic and the inherent innocence of childhood desserts serves as a powerful "
        "psychological differentiator in a crowded farmers\u2019 market environment, appealing to "
        "modern consumers\u2019 appreciation for irony, experiential retail, and high-quality "
        "ingredients."
    )

    # ── SECTION 3: FARMERS MARKET ECOSYSTEM ───────────────────
    add_heading_styled(doc, "3. The Santa Monica Farmers Market Ecosystem & RFP Dynamics", level=1)

    add_heading_styled(doc, "3.1 The Request for Proposal (RFP) Barrier", level=2)

    add_body_para(doc,
        "Operating within the municipal boundaries of Santa Monica, and specifically within its "
        "highly curated farmers\u2019 markets, requires navigating a complex matrix of overlapping "
        "jurisdictions. The operational hurdle that dictates the timeline for the entire "
        "Armageddon Treats business plan is securing market access."
    )

    add_body_para(doc,
        "The Santa Monica Farmers Markets are not open-enrollment platforms where vendors can "
        "simply pay a fee to secure a stall. Instead, the municipality tightly controls the "
        "vendor mix to maintain the market\u2019s reputation for quality, diversity, and "
        "agricultural integrity. Vendors are selected exclusively through a formalized City bid "
        "process, organized as a Request for Proposals (RFP)."
    )

    add_body_para(doc,
        "The application window for this RFP is notoriously narrow. Applications are strictly "
        "closed for the 2025 and 2026 cycles outside of the designated bid periods. The "
        "application process for the Pico and Main Street markets typically opens only once per "
        "year, occurring in March. Because the application window is so brief, the business must "
        "possess all requisite city and county health permits, commissary agreements, and "
        "finalized business plans prior to submission to be considered a viable, ready-to-operate "
        "candidate by the selection committee. A failure to synchronize the operational launch "
        "with the March application window results in a mandatory twelve-month delay in accessing "
        "the primary intended revenue stream."
    )

    add_heading_styled(doc, "3.2 Vendor Classification and Variable Fee Structures", level=2)

    add_body_para(doc,
        "Within the market ecosystem, vendors are segregated into distinct classifications with "
        "varying financial obligations."
    )

    add_styled_table(doc,
        ["Vendor Type", "Fee Structure", "Notes"],
        [
            ["Agricultural (CPCs)", "6% of gross daily sales + $2 state fee", "Farmers, ranchers, fishers"],
            ["Prepared & Pre-packaged Food", "12% \u2013 13% of gross daily sales", "Armageddon Treats classification"],
        ],
        col_widths=[2.0, 2.3, 2.5]
    )

    add_body_para(doc,
        "While a 12\u201313% top-line revenue sacrifice is substantial, this variable fee "
        "structure is actually highly advantageous for a thinly capitalized startup. Unlike "
        "traditional commercial real estate leases or ghost kitchen rentals that demand fixed "
        "monthly overhead regardless of performance, the farmers\u2019 market model converts real "
        "estate costs into a purely variable expense. If inclement weather or seasonal shifts "
        "result in lower foot traffic, the vendor\u2019s financial burden scales down "
        "proportionately, significantly lowering the enterprise\u2019s break-even point and "
        "mitigating the risk of insolvency during the critical first year of operations."
    )

    add_heading_styled(doc, "3.3 Market Attendance and Operational Rigor", level=2)

    add_body_para(doc,
        "The City of Santa Monica enforces strict attendance and operational policies to ensure "
        "the markets remain vibrant and fully populated. Prepared and packaged food vendors "
        "operating under a signed permit are allowed only three excused absences per calendar "
        "year for any reason, including illness, staffing shortages, or weather conditions. "
        "During these three excused absences, the standard market fees are waived. However, "
        "beyond these three allowances, vendors are required to pay their average historical "
        "market fees regardless of the reason for their absence."
    )

    add_body_para(doc,
        "This policy demands robust operational reliability from the vendor; a failure in the "
        "tow vehicle or a disruption in the supply chain that prevents attendance directly "
        "penalizes the business\u2019s bottom line. Furthermore, the market enforces strict "
        "sanitary boundaries; no live animals are permitted within twenty feet of any area where "
        "food is stored or prepared, necessitating careful site management by the vendor."
    )

    # ── SECTION 4: MUNICIPAL COMPLIANCE ───────────────────────
    add_heading_styled(doc, "4. Municipal Compliance & Sustainability Mandates", level=1)

    add_heading_styled(doc, "4.1 Licensing Fees and Tax Structures", level=2)

    add_body_para(doc,
        "Operating legally within Santa Monica limits requires obtaining a specific Seller\u2019s "
        "Permit (issued free of charge from the California Department of Tax and Fee "
        "Administration) and a Santa Monica Business License combined with a localized Vendor "
        "Permit."
    )

    add_body_para(doc,
        "The base processing and regulatory cost for a vendor permit is approximately $157.96 to "
        "$175.60 per year, which includes regulatory permit fees and state-mandated ADA "
        "compliance (CASp) fees. Additionally, under the municipal tax code\u2019s Tax Rate "
        "Group IX, vehicle and sidewalk vending operations are required to pay a base tax of "
        "$50 per vehicle or vending operation annually, in addition to the vendor permit costs. "
        "To obtain these licenses, the vendor must present a valid Los Angeles County Department "
        "of Public Health permit, proving that the sequence of regulatory approvals must begin "
        "at the county level before municipal licenses can be finalized."
    )

    add_heading_styled(doc, "4.2 The Disposable Food Service Ware Ordinance", level=2)

    add_body_para(doc,
        "Santa Monica enforces a comprehensive Disposable Food Service Ware Ordinance that "
        "applies to any business, organization, or individual providing prepared food or "
        "beverages for public consumption within the city. Under this ordinance, the use of all "
        "expanded polystyrene (commonly known as Styrofoam), whether foam or clear plastic #6, "
        "is strictly banned across all food establishments and mobile vendors. All prepared "
        "food must be served in marine-degradable, compostable, or highly recyclable packaging."
    )

    add_body_para(doc,
        "For a shaved ice and cotton candy vendor, this necessitates the outright rejection of "
        "cheap, conventional plastic serving vessels in favor of Polylactic Acid (PLA) "
        "bioplastics or sugar cane pulp alternatives."
    )

    add_heading_styled(doc, "4.3 Green Business Certification Alignment", level=2)

    add_body_para(doc,
        "While official \u201cGreen Business Certification\u201d (GBC) in Santa Monica requires "
        "a business to occupy a minimum of 500 square feet of commercial space (excluding mobile "
        "carts and home offices), adhering to the program\u2019s underlying principles is a "
        "strategic necessity for Armageddon Treats. Integrating compostable PLA plastics, "
        "minimizing water waste, and maintaining a verifiable organic supply chain will heavily "
        "favor Armageddon Treats during the competitive Farmers Market RFP grading process. The "
        "selection committees prioritize vendors whose operational models seamlessly align with "
        "the city\u2019s broader ecological goals."
    )

    # ── SECTION 5: LA COUNTY HEALTH REGS ─────────────────────
    add_heading_styled(doc, "5. Los Angeles County Public Health Regulations", level=1)

    add_heading_styled(doc, "5.1 Risk Classifications and Permit Fees", level=2)

    add_body_para(doc,
        "Food vending in Los Angeles County is intensely regulated to prevent foodborne illness "
        "and ensure public safety. Recent legislative shifts at the state level, particularly "
        "the passage of SB 972 in 2022 (effective January 2023), sought to decriminalize and "
        "regulate sidewalk vending by introducing the Compact Mobile Food Operation (CMFO) "
        "classification into the California Retail Food Code (CRFC). However, because the "
        "Armageddon Treats business model involves active food preparation\u2014spinning heated "
        "sugar, shaving bulk ice, and dispensing liquid syrups\u2014it triggers significantly "
        "more stringent health code requirements."
    )

    add_styled_table(doc,
        ["Facility Classification", "Description", "FY 2025\u20132026 Annual Permit Fee"],
        [
            ["CMFO, Low Risk", "Prepackaged non-hazardous foods, whole produce", "$126.00"],
            ["CMFO, Moderate Risk", "Limited food preparation (e.g., shaved ice carts)", "$299.00"],
            ["MFF, Low Risk", "Prepackaged ice cream trucks, produce trucks", "$325.00"],
            ["MFF, Moderate Risk", "Soft serve, smoothie, or beverage trucks/trailers", "$598.00"],
        ],
        col_widths=[1.8, 3.0, 2.0]
    )

    add_body_para(doc,
        "For Armageddon Treats, assuming the operation utilizes a towable trailer rather than a "
        "pushcart, it will likely be classified as a Moderate Risk Mobile Food Facility, "
        "incurring an annual public health permit fee of $598. Before this permit is issued, the "
        "physical vehicle must undergo a comprehensive plan check to ensure it meets structural "
        "health codes. The plan check fee adds an additional minimum cost of $347, which scales "
        "depending on the complexity of the vehicle\u2019s design."
    )

    add_heading_styled(doc, "5.2 Mandatory Equipment and Structural Specifications", level=2)

    add_body_para(doc,
        "The Los Angeles County Department of Public Health (LACDPH) enforces rigid structural "
        "mandates for mobile food facilities to ensure sanitary conditions during field operations."
    )

    # Equipment specs table
    add_styled_table(doc,
        ["Requirement", "Specification"],
        [
            ["Equipment Certification", "All food-contact equipment must be ANSI-accredited (NSF certified); electrical appliances must meet UL standards"],
            ["Handwashing Sink", "Minimum 9\u2033W \u00d7 9\u2033L \u00d7 5\u2033D; warm water (100\u2013108\u00b0F) via pressurized or gravity-fed system"],
            ["Warewashing (3-Compartment Sink)", "Each compartment must submerge largest utensil; water temperature \u2265 120\u00b0F for thermal sanitization; two integral metal drainboards"],
            ["Refrigeration", "Minimum 12 cu. ft. usable space for perishable ingredients"],
            ["Enclosure", "Solid walls or 16-mesh screens; service windows limited to 216 sq. in. maximum"],
        ],
        col_widths=[2.2, 4.6]
    )

    add_heading_styled(doc, "5.3 The Commissary Kitchen Imperative", level=2)

    add_body_para(doc,
        "A foundational, non-negotiable requirement of the Los Angeles County health code is "
        "that all mobile food facilities\u2014regardless of size\u2014must operate in conjunction "
        "with an approved, permanent commissary kitchen or Mobile Support Unit (MSU). The storage "
        "of a mobile food facility, or the storage of any food products, packaging, or supplies "
        "associated with the business, at a private home or residential garage is explicitly "
        "prohibited by law."
    )

    add_body_para(doc,
        "The commissary serves as the hygienic operational hub for the business. It is the "
        "designated location where raw ingredients are safely stored, deep cleaning of the mobile "
        "unit occurs using industrial floor drains, fresh potable water tanks are filled, and "
        "liquid greywater and grease are legally and safely disposed of into municipal sanitary "
        "sewer systems."
    )

    add_heading_styled(doc, "5.4 Commissary Logistics and Market Rates", level=2)

    add_body_para(doc,
        "Securing a commissary contract is a prerequisite for obtaining the county health permit. "
        "The owner must submit a signed \u201cVerification of Proper Food Vehicle Storage\u201d "
        "form from the commissary operator to the LACDPH. To ensure ongoing compliance, health "
        "inspectors conduct random audits to verify that mobile units are actively utilizing "
        "their assigned commissaries and maintaining an approved \u201cRoute Sheet\u201d that "
        "details the vehicle\u2019s daily operating locations."
    )

    add_styled_table(doc,
        ["Commissary Model", "Cost Range", "Best For"],
        [
            ["Hourly Shared-Use Kitchen", "$20 \u2013 $120 / hour (avg. $25\u2013$65)", "Occasional prep; low-volume operators"],
            ["Monthly Lease (< 800 sq ft shared)", "$1,200 \u2013 $2,500 / month", "Full-service operators needing dedicated space"],
            ["Parking + Wash-Down Only", "$300 \u2013 $1,500 / month", "Mobile vendors requiring storage, dumping, and water fill"],
        ],
        col_widths=[2.2, 2.3, 2.3]
    )

    # ── SECTION 6: MOBILE ASSET PROCUREMENT ──────────────────
    add_heading_styled(doc, "6. Mobile Asset Procurement Strategy", level=1)

    add_heading_styled(doc, "6.1 The Financial Impossibility of a Motorized Truck", level=2)

    add_body_para(doc,
        "A food truck is a fully self-contained vehicle integrating the engine, drivetrain, and "
        "commercial kitchen into a single chassis. While they offer superior mobility and a "
        "smaller parking footprint, they are capital-intensive. In the 2025\u20132026 Southern "
        "California secondary market, even heavily used, older-model food trucks command prices "
        "between $40,000 and $80,000. New or custom-built motorized units routinely exceed "
        "$100,000 to $150,000. Therefore, procuring a functional, health-code-compliant "
        "motorized food truck is mathematically impossible within the constraints of a $14,000 "
        "SBA microloan."
    )

    add_heading_styled(doc, "6.2 The Concession Trailer Pathway", level=2)

    add_body_para(doc,
        "The only viable pathway to market is the acquisition of a towable concession trailer. "
        "A trailer relies on a separate towing vehicle for transportation but provides an "
        "identical, compliant kitchen environment once parked."
    )

    add_styled_table(doc,
        ["Procurement Route", "Size", "Price Range", "Advantages"],
        [
            ["Used Secondary Market", "10\u201316 ft", "$12,000 \u2013 $25,000", "Immediate availability; may include equipment"],
            ["Direct-to-Manufacturer (Import)", "\u224813 ft (400 cm)", "$7,000 \u2013 $8,000 + freight", "New asset; customizable; brand-ready for matte-black wrap"],
        ],
        col_widths=[2.0, 1.2, 1.8, 1.8]
    )

    add_heading_styled(doc, "6.3 The Tow Vehicle Dependency", level=2)

    add_body_para(doc,
        "The critical caveat to the trailer strategy is the requirement of a robust towing "
        "vehicle. A fully loaded 10-foot concession trailer, carrying commercial equipment, "
        "50 gallons of fresh water (approximately 415 pounds), and generators, requires a "
        "vehicle with a substantial towing capacity. Capable half-ton trucks (e.g., 2010\u20132015 "
        "Ford F-150 or Chevrolet Silverado 1500) carry an average procurement cost of $12,000 "
        "to $18,000 in Los Angeles."
    )

    add_body_para(doc,
        "Because the $14,000 SBA loan cannot cover the cost of both the concession trailer and "
        "the required tow vehicle, the operational viability of Armageddon Treats rests on an "
        "absolute prerequisite: The veteran proprietor must already personally own a capable "
        "towing vehicle. If a tow vehicle must be acquired to launch the business, the $14,000 "
        "capitalization is fatally insufficient, and operations cannot commence without securing "
        "substantial secondary financing.",
        bold=False, italic=True
    )

    # ── SECTION 7: UNIT ECONOMICS ────────────────────────────
    add_heading_styled(doc, "7. Micro-Economic Analysis: Product Unit Economics", level=1)

    add_body_para(doc,
        "The strategic brilliance of focusing a mobile food operation on shaved ice and cotton "
        "candy lies in the unparalleled unit economics of the products. These specific desserts "
        "boast some of the highest gross margins in the entire food service industry."
    )

    add_heading_styled(doc, "7.1 Organic Cotton Candy", level=2)

    add_body_para(doc,
        "Organic floss sugar\u2014available in sophisticated flavors such as strawberry, "
        "watermelon, and almond\u2014is formulated without artificial additives and provides a "
        "cleaner taste profile. The wholesale cost is approximately $18.73 per 3.25-pound "
        "carton, yielding 50\u201360 standard servings."
    )

    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Raw Material Cost per Serving", "$0.31 \u2013 $0.37"],
            ["Retail Price Point", "$5.00 \u2013 $6.00"],
            ["Gross Profit Margin", "\u224894%"],
            ["Net Profit per Unit (after indirect costs)", "$3.50 \u2013 $4.00"],
        ],
        col_widths=[3.4, 3.4]
    )

    add_heading_styled(doc, "7.2 Premium Shaved Ice", level=2)

    add_body_para(doc,
        "To comply with Santa Monica\u2019s Disposable Food Service Ware Ordinance, traditional "
        "plastic cups must be replaced with biodegradable PLA flower cups. Wholesale pricing for "
        "1,200-piece cases of 12 oz PLA flower cups is approximately $300 (unit cost: $0.25). "
        "When factoring in high-quality, pure cane sugar syrups and the ice itself, the total "
        "COGS for a premium shaved ice dessert remains under $0.60."
    )

    add_heading_styled(doc, "7.3 Revenue Projections", level=2)

    add_styled_table(doc,
        ["Scenario", "Units / Day", "Daily Gross Revenue", "Monthly Gross (Weekends Only)"],
        [
            ["Conservative", "100 units @ $4.00 avg.", "$400", "$3,200"],
            ["Optimized", "200 units @ $5.00 avg.", "$1,000", "$8,000"],
        ],
        col_widths=[1.5, 2.0, 1.7, 1.6]
    )

    # ── SECTION 8: FINANCIAL ARCHITECTURE ────────────────────
    add_heading_styled(doc, "8. Financial Architecture: SBA 7(a) Microloan & Capital Allocation", level=1)

    add_heading_styled(doc, "8.1 Loan Parameters and Interest Rate Mechanics", level=2)

    add_body_para(doc,
        "The primary financing mechanism is a $14,000 microloan procured through the SBA\u2019s "
        "7(a) loan program. The SBA does not lend money directly; instead, it partners with "
        "approved lenders and guarantees up to 85% for loans under $150,000, thereby reducing "
        "the lender\u2019s risk and increasing access to capital."
    )

    add_body_para(doc,
        "For 7(a) loans of $50,000 or less, the maximum interest rate is capped at the base "
        "prime rate plus 6.5%. With the current prime rate at 6.75% (March 2026), the absolute "
        "maximum applicable interest rate for this microloan is 13.25%. Extended maturity terms "
        "of up to 10 years (120 months) ensure highly manageable monthly debt service payments."
    )

    add_heading_styled(doc, "8.2 The Veteran Advantage and Alternative Funding", level=2)

    add_body_para(doc,
        "Under the SBA Veterans Advantage initiative, the SBA completely waives the upfront "
        "guarantee fee for qualifying veterans (honorable discharge, \u226551% ownership) on "
        "7(a) loans under $150,000. This waiver eliminates a significant closing cost, allowing "
        "more of the $14,000 principal to be deployed directly into physical assets."
    )

    add_body_para(doc,
        "The proprietor must also actively pursue non-dilutive grant funding to bolster working "
        "capital reserves:"
    )

    add_styled_table(doc,
        ["Grant / Program", "Value", "Eligibility Notes"],
        [
            ["Stephen L. Tadlock Veteran Business Grant", "$1,000", "Veteran-owned small businesses"],
            ["Melissa Washington Small Business Award", "Up to $5,000", "Eligible veteran ventures"],
            ["DVBE / LSBE Certification (LA County)", "Up to 15% bid preference", "Municipal catering & event contract bidding advantage"],
        ],
        col_widths=[2.8, 1.5, 2.5]
    )

    add_heading_styled(doc, "8.3 Strict Capital Allocation Strategy", level=2)

    add_body_para(doc,
        "Allocating a mere $14,000 to legally launch a compliant mobile food facility in "
        "Los Angeles County leaves absolutely zero margin for error. The capital must be "
        "deployed with extreme precision."
    )

    add_styled_table(doc,
        ["Expense Category", "Allocation", "Strategic Justification"],
        [
            ["Concession Trailer", "$7,500", "Compliant 10 ft imported trailer with basic electrical and multi-compartment sinks"],
            ["Commercial Equipment & Branding", "$2,000", "ANSI/NSF certified shaved ice machine, cotton candy spinner, matte-black vinyl wrap"],
            ["Permits & Health Licensing", "$1,200", "LACDPH plan check, MFF/CMFO annual fee, Santa Monica Vendor License"],
            ["Initial Insurance Premiums", "$1,000", "Down payments on Commercial General Liability and Commercial Auto policies"],
            ["Initial Inventory & Packaging", "$800", "Bulk organic floss sugars, pure cane syrups, biodegradable PLA flower cups"],
            ["Commissary Deposit / Rent", "$500", "First month\u2019s rent or hourly block to secure LACDPH storage verification"],
            ["Working Capital / Contingency", "$1,000", "Cash reserve for unforeseen compliance modifications, fuel, or delays"],
            ["TOTAL CAPITAL DEPLOYMENT", "$14,000", "Fully exhausted SBA 7(a) Microloan"],
        ],
        col_widths=[2.2, 1.2, 3.4]
    )

    # ── SECTION 9: OPEX & LABOR ──────────────────────────────
    add_heading_styled(doc, "9. Operational Expenses (OPEX) & Labor Cost Dynamics", level=1)

    add_heading_styled(doc, "9.1 Insurance and Liability Management", level=2)

    add_body_para(doc,
        "Operating a mobile food unit requires robust, multi-tiered insurance coverage. "
        "Commercial General Liability insurance averages approximately $120 per month in "
        "California. Commercial Auto Insurance for food transport applications ranges from "
        "$156 to over $260 per month, representing one of the largest fixed monthly overhead "
        "costs for the operation."
    )

    add_heading_styled(doc, "9.2 Fuel, Energy, and Maintenance", level=2)

    add_body_para(doc,
        "Food trucks and heavy tow vehicles typically average a dismal 7 to 10 miles per gallon "
        "during urban operation. Assuming regular weekend travel across the Los Angeles "
        "metropolitan area, monthly fuel costs are projected at $250\u2013$500. If market "
        "locations do not provide direct shore power, portable generators introduce a secondary "
        "fuel cost (propane or diesel)."
    )

    add_heading_styled(doc, "9.3 Labor Regulations and Minimum Wage Escalation", level=2)

    add_body_para(doc,
        "Initially, the business will likely operate as a sole proprietorship driven by the "
        "sweat equity of the veteran founder. As the business scales, hiring part-time assistants "
        "will become a necessity."
    )

    add_styled_table(doc,
        ["Jurisdiction", "Minimum Wage (2026)", "Effective Date"],
        [
            ["California (Statewide)", "$16.90 / hour", "January 1, 2026"],
            ["City of Los Angeles", "$18.42 / hour", "July 1, 2026"],
            ["Unincorporated LA County", "$18.47 / hour", "July 1, 2026"],
        ],
        col_widths=[2.5, 2.0, 2.3]
    )

    add_body_para(doc,
        "When factoring in mandatory employer-side payroll taxes, unemployment insurance, and "
        "workers\u2019 compensation premiums ($78\u2013$91/month for food vendors), the true "
        "burdened cost of a part-time service assistant approaches $22.00\u2013$24.00 per hour."
    )

    # Monthly OPEX Summary
    add_heading_styled(doc, "9.4 Estimated Monthly OPEX Summary", level=2)

    add_styled_table(doc,
        ["Expense", "Low Estimate", "High Estimate"],
        [
            ["Commercial General Liability Insurance", "$120", "$120"],
            ["Commercial Auto Insurance", "$156", "$260"],
            ["Fuel (Tow Vehicle)", "$250", "$500"],
            ["Generator Fuel (Propane/Diesel)", "$50", "$150"],
            ["Commissary Rent", "$300", "$1,500"],
            ["Stall Fees (12\u201313% of gross)", "$384", "$1,040"],
            ["Part-Time Labor (if applicable)", "$0", "$768"],
            ["Miscellaneous / Maintenance", "$100", "$200"],
            ["TOTAL ESTIMATED MONTHLY OPEX", "$1,360", "$4,538"],
        ],
        col_widths=[3.2, 1.7, 1.9]
    )

    # ── SECTION 10: STRATEGIC CONCLUSIONS ────────────────────
    add_heading_styled(doc, "10. Strategic Conclusions & Risk Mitigation", level=1)

    add_body_para(doc,
        "The \u201cArmageddon Treats\u201d business plan presents a fascinating study in "
        "operational extremes. It combines a highly provocative, culturally relevant branding "
        "strategy with core products that boast virtually unmatched profitability in the food "
        "service sector. The thematic positioning of an apocalyptic, \u201cdoomsday\u201d "
        "aesthetic combined with sustainable, USDA organic ingredients aligns seamlessly with "
        "the complex, experiential desires of the affluent Santa Monica consumer demographic."
    )

    add_body_para(doc,
        "However, executing this vision on a heavily constrained $14,000 SBA 7(a) microloan "
        "demands flawless strategic foresight and unyielding financial discipline. There is "
        "absolutely zero tolerance for extraneous expenditures or regulatory missteps."
    )

    add_body_para(doc, "The success of the enterprise hinges on the execution of a sequential, "
        "highly coordinated rollout:", bold=True
    )

    # Key imperatives as a table
    add_styled_table(doc,
        ["Priority", "Imperative", "Detail"],
        [
            ["1", "Asset Sequencing",
             "Procure a compliant, low-cost towable asset (relying on pre-existing tow vehicle ownership) "
             "and secure a foundational agreement with a licensed commercial commissary kitchen."],
            ["2", "Regulatory Timing",
             "Finalize all LA County Health permits and vehicle plan checks prior to the brief Santa Monica "
             "Farmers Market RFP application window in March. Missing this window forces a 12-month delay."],
            ["3", "Capital Preservation",
             "Aggressively leverage veteran-owned status for SBA fee waivers and pursue supplementary "
             "state-level grant funding to build a defensive working capital reserve against volatile "
             "insurance and fuel costs."],
        ],
        col_widths=[0.6, 1.6, 4.6]
    )

    add_body_para(doc,
        "By adhering strictly to municipal sustainability ordinances, executing a disciplined "
        "capital allocation strategy, and leaning into the high-margin mathematics of its core "
        "dessert offerings, Armageddon Treats possesses the structural foundation required to "
        "transition from a theoretical concept into a profitable, highly visible fixture within "
        "Southern California\u2019s premier outdoor retail ecosystem."
    )

    # Note about Works Cited
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("For full source documentation, see the companion document: "
                     "Armageddon_Treats_Works_Cited.docx")
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run.font.color.rgb = MID_GRAY
    run.italic = True

    doc.save("output/Armageddon_Treats_Business_Plan_Audit.docx")
    print("Main document saved: output/Armageddon_Treats_Business_Plan_Audit.docx")


if __name__ == "__main__":
    build_main_document()
