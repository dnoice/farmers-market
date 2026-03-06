"""
Generate the Armageddon Treats Works Cited / References document.
"""

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor

from armageddon_treats.generators.audit import (
    set_cell_shading, set_cell_borders, style_header_row, style_data_row,
    add_styled_table, add_heading_styled, add_body_para,
    setup_document, add_header_footer,
    BLACK, DARK_GRAY, MID_GRAY, ACCENT, WHITE,
    FONT_BODY, FONT_HEADING,
)


# ══════════════════════════════════════════════════════════════
#  WORKS CITED DOCUMENT
# ══════════════════════════════════════════════════════════════
def build_works_cited():
    doc = setup_document()
    add_header_footer(doc, "Armageddon Treats — Works Cited")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("ARMAGEDDON TREATS")
    run.font.name = FONT_HEADING
    run.font.size = Pt(24)
    run.font.color.rgb = BLACK
    run.bold = True
    run.font.small_caps = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run("Works Cited")
    run2.font.name = FONT_HEADING
    run2.font.size = Pt(16)
    run2.font.color.rgb = MID_GRAY

    # Divider
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(16)
    run3 = p3.add_run("_" * 60)
    run3.font.color.rgb = ACCENT
    run3.font.size = Pt(10)

    # The citations
    citations = [
        ("Data and Research", "Downtown Santa Monica", "https://downtownsm.com/data-and-research"),
        ("The Santa Monica Farmer's Market Rolls On", "Media Milwaukee", "https://mediamilwaukee.com/features/santa-monica-farmers-market-downtown/"),
        ("Farmers Market", "santamonica.gov", "https://www.santamonica.gov/categories/programs/farmers-market"),
        ("Certified Farmers' Markets by County as of January 1, 2026", "California Department of Food and Agriculture", "https://www.cdfa.ca.gov/is/docs/CurrentMrktsCounty.pdf"),
        ("Rules, Regulations and Policy Guide", "City of Santa Monica", "https://www.santamonica.gov/media/Farmers_Market/Farmers%20Market%20Rules/Farmers%20Market%20Rules%202021.pdf"),
        ("Santa Monica Sunday Main Street Farmers Markets", "Foresyte", "https://foresyteapp.com/santa-monica-sunday-main-street-farmers-markets/"),
        ("Mythologized Counter-Futures and Self-Protective Consumption: A Netnography of Doomsday Preppers", "Journal of Consumer Research, Oxford Academic", "https://academic.oup.com/jcr/article/52/4/759/7991296"),
        ("Believing in the End: An Ethnography of the Prepper Culture", "Pennsylvania State University, Electronic Theses and Dissertations", "https://etda.libraries.psu.edu/files/final_submissions/18481"),
        ("Prepping as Implicit Activism: Risk, Danger, and Post-Capitalist Imaginaries in Prepper Literature", "Taylor & Francis", "https://www.tandfonline.com/doi/full/10.1080/14742837.2024.2349568"),
        ("How We Consume Fear in Times of Crisis, Brands That Change the Story", "Concept Bureau", "https://conceptbureau.com/"),
        ("Corporate Marketing: Apocalypse, Advent and Epiphany", "ResearchGate", "https://www.researchgate.net/publication/242342716_Corporate_marketing_Apocalypse_advent_and_epiphany"),
        ("From 'Aesthetic' to Aestheticization: A Multi-Layered Cultural Approach", "Taylor & Francis", "https://www.tandfonline.com/doi/full/10.1080/10253866.2021.1935900"),
        ("Packaging of Ready-to-Eat Products: A Mix Research Approach for Exploring Aesthetic Element Influence on Brand Preference", "British Food Journal, Emerald Publishing", "https://www.emerald.com/bfj/article/125/11/4182/65748/"),
        ("Visual Design Cues Impacting Food Choice: A Review and Future Research Agenda", "PMC / National Library of Medicine", "https://pmc.ncbi.nlm.nih.gov/articles/PMC7589873/"),
        ("Mind the Gap", "McKinsey & Company", "https://www.mckinsey.com/featured-insights/mind-the-gap"),
        ("Color in Healthcare Environments", "The Center for Health Design", "https://www.healthdesign.org/sites/default/files/color_in_hc_environ.pdf"),
        ("Apply to be a Food Vendor at the Farmers Market", "santamonica.gov", "https://www.santamonica.gov/process-explainers/apply-to-be-a-food-vendor-at-the-farmers-market"),
        ("Sell At A Market", "California Farmers' Markets Association", "http://www.cafarmersmkts.com/sell-at-a-farmers-market"),
        ("Apply to be a Vendor at the Santa Monica Farmers Market", "santamonica.gov", "https://www.santamonica.gov/process-explainers/apply-to-be-a-vendor-at-the-santa-monica-farmers-market"),
        ("How to Become a Licensed Vendor in Santa Monica", "santamonica.gov", "https://www.santamonica.gov/blog/how-to-become-a-licensed-vendor-in-santa-monica"),
        ("For Our Vendors", "Food Access LA", "https://foodaccessla.org/vendors"),
        ("How to Get a Vehicle Vending Permit", "City of Santa Monica", "https://www.santamonica.gov/process-explainers/how-to-get-a-vehicle-vending-permit"),
        ("Schedule of License Fees by Business Classification and Tax Rate Group", "City of Santa Monica, eCode360", "https://ecode360.com/42734989"),
        ("Vending", "City of Santa Monica, eCode360", "https://ecode360.com/42735317"),
        ("Vending Safely in Santa Monica", "santamonica.gov", "https://www.santamonica.gov/topic-explainers/vending-in-santa-monica"),
        ("Disposable Food Service Ware Ordinance", "santamonica.gov", "https://www.santamonica.gov/disposable-food-service-ware-ordinance"),
        ("Santa Monica Green Business Certification Checklist", "SMGBC", "http://smgbc.org/"),
        ("Green Business Certification", "santamonica.gov", "https://www.santamonica.gov/programs/green-business-certification"),
        ("How to Become a Certified Green Business", "City of Santa Monica", "https://www.santamonica.gov/process-explainers/how-to-become-a-certified-green-business"),
        ("Santa Monica Green Business Certification Application Checklist for Restaurants & Grocers", "Responsible Purchasing Network", "https://www.responsiblepurchasing.org/"),
        ("Selective Sourcing", "Santa Monica Daily Press", "https://www.smdp.com/152165-2/"),
        ("Compact Mobile Food Operation in Los Angeles County", "LA County Dept. of Public Health", "http://publichealth.lacounty.gov/eh/business/compact-mobile-food-operation.htm"),
        ("FY 2025\u20132026 Public Health and Financial Management Fee Schedule", "County of Los Angeles", "http://publichealth.lacounty.gov/eh/docs/permit/public-health-financial-management-fee-schedule.pdf"),
        ("New Mobile Food Facility Fees Summary", "County of Los Angeles Public Health", "http://publichealth.lacounty.gov/eh/docs/permit/new-mobile-food-facility-fees-summary.pdf"),
        ("How to Get a Health Permit for Your Kitchen in Los Angeles", "CloudKitchens", "https://cloudkitchens.com/blog/how-to-get-a-health-permit-for-restaurants-in-los-angeles/"),
        ("Starter Kit for Mobile Food Vending", "LA Business Navigator, City of Los Angeles", "https://business.lacity.gov/plan-business/starter-kits/starter-kit-mobile-food-vending"),
        ("Mobile Food Facility Plan Check Guideline", "LA County Dept. of Public Health", "http://publichealth.lacounty.gov/eh/inspection/plan-check-operational-guidelines-mobile-food-facilities-mobile-support-unit.htm"),
        ("Health Code Requirements: Community Events in LA County", "Santa Monica College", "https://www.smc.edu/community/office-of-student-life/student-government/documents/AS-State-Forms-and-Docs/ca-la-county-health-code.pdf"),
        ("Food Service Trucks, Trailers, and Boats", "LA County Dept. of Public Health", "http://publichealth.lacounty.gov/eh/business/food-trucks-carts.htm"),
        ("Your Guide to Food Truck Commissary Requirements", "The Restaurant Warehouse", "https://therestaurantwarehouse.com/blogs/restaurant-equipment/food-truck-commissary-requirements"),
        ("Commissary Guidance", "LA County Dept. of Public Health", "http://publichealth.lacounty.gov/eh/docs/permit/commissary-guidance-facility-owners.pdf"),
        ("How Much Does it Cost to Rent a Commercial Kitchen? (2024 Costs)", "Toast POS", "https://pos.toasttab.com/blog/on-the-line/how-much-does-it-cost-to-rent-a-commercial-kitchen"),
        ("Renting a Commercial Kitchen: Cost Breakdown [2025]", "Mobile Culinaire", "https://www.mobileculinaire.com/resource-center/commercial-kitchen-rent-cost-2025"),
        ("Food Truck Commissary Costs and Rentals: A Guide for Food Entrepreneurs", "Le Gourmet Factory", "https://www.legourmetfactory.com/blogs/news/food-truck-commissary-costs-and-rentals"),
        ("Food Truck vs Food Trailer: Which Is More Profitable?", "Kinema Trailer Depot", "https://kinematrailerdepot.com/food-truck-vs-food-trailer/"),
        ("How Much Does a Food Truck Cost? Complete Pricing Guide", "CloudKitchens", "https://cloudkitchens.com/blog/how-much-does-a-food-truck-cost/"),
        ("Used Concession Trailer For Sale Near Orange County, CA", "Commercial Truck Trader", "https://www.commercialtrucktrader.com/"),
        ("Concession Trailer For Sale Near Los Angeles, CA", "Commercial Truck Trader", "https://www.commercialtrucktrader.com/"),
        ("Food Trailers for Sale in Los Angeles, CA", "Oulead Trailer", "https://www.foodtrailersale.com/"),
        ("Used Ford F-150 Trucks for Sale Near Los Angeles, CA Under $20,000", "Cars.com", "https://www.cars.com/"),
        ("Used 2010 Chevrolet Silverado 1500 for Sale in Los Angeles, CA", "Edmunds", "https://www.edmunds.com/"),
        ("Used Ford F-150 for Sale near Los Angeles, CA", "CarGurus", "https://www.cargurus.com/"),
        ("Organic Cotton Candy Sugar", "Nature's Flavors", "https://www.naturesflavors.com/collections/organic-cotton-candy-sugar"),
        ("Organic Almond Cotton Candy Floss Sugar | Vegan & Gluten-Free", "Best Flavors", "https://www.bestflavors.com/products/almond-cotton-candy-floss-sugar-organic"),
        ("Cotton Candy Supplies", "Popcorn Supply", "https://popcornsupply.com/cotton-candy-supplies"),
        ("How Much Do Cotton Candy Vendors Make? A Complete Profit Breakdown", "Wider Matrix", "https://widermatrix.com/how-much-do-cotton-candy-vendors-make/"),
        ("Flower Cups \u2014 4-ounce, 6-ounce, & 8-ounce sizes", "Hawaiian Shaved Ice", "https://hawaiianshavedice.com/products/flower-cups"),
        ("Eco Friendly Cups", "Shave Ice Supplies", "https://shaveicesupplies.com/eco-friendly-ice-cream-cups/"),
        ("Flower Shaved Ice Cups", "Real Hawaiian Ice", "https://realhawaiianice.com/product-category/flower-cups/"),
        ("How to Start a Snow Cone Business & Make It Profitable", "Webstaurant Store", "https://www.webstaurantstore.com/article/725/snow-cone-business.html"),
        ("How Much Can You Make with a Shave Ice Business?", "Sno Biz", "https://www.snobiz.com/how-much-can-you-make-with-a-shave-ice-business/"),
        ("Profitability of a Shave Ice Business?", "IcySkyy", "https://www.icy-sky.com/blogs/news/profitability-of-a-shave-ice-business"),
        ("7(a) Loans", "U.S. Small Business Administration", "https://www.sba.gov/funding-programs/loans/7a-loans"),
        ("Loans", "U.S. Small Business Administration", "https://www.sba.gov/funding-programs/loans"),
        ("SBA Loan Rates 2026", "NerdWallet", "https://www.nerdwallet.com/business/loans/learn/sba-loan-rates"),
        ("Business Loans for Veterans in 2026", "VA Loan Network", "https://valoannetwork.com/business-loans-for-veterans/"),
        ("Veteran Contracting Assistance Programs", "U.S. Small Business Administration", "https://www.sba.gov/federal-contracting/contracting-assistance-programs/veteran-contracting-assistance-programs"),
        ("SBA Loans and Small Business Loans for Veterans 2026", "Your Funding Tree", "https://www.yourfundingtree.com/exploring-sba-loans-for-veterans/"),
        ("Veterans Advantage Program: Exclusive SBA Benefits for Service Members", "Lendio", "https://www.lendio.com/blog/sba-veterans-advantage-program"),
        ("Grants for Veteran-owned Businesses", "SoCal VBOC", "https://socalvboc.org/grants/"),
        ("Office of Small Business \u2014 Certifications", "LA County", "https://www.ajcc.lacounty.gov/businesses/office-of-small-business/certifications"),
        ("Apply for or Re-Apply for Certification as a Small Business and/or Disabled Veteran Business Enterprise", "DGS.ca.gov", "https://www.dgs.ca.gov/PD/Services/Page-Content/Procurement-Division-Services-List-Folder/Apply-for-or-Re-apply-as-Small-Business-Disabled-Veteran-Business-Enterprise"),
        ("Small Business & Disabled Veterans Business Enterprise (SB & DVBE)", "CalPERS", "https://www.calpers.ca.gov/about/doing-business-with-calpers/small-business-dvbe"),
        ("Food Business Insurance Cost (2026 Rates)", "MoneyGeek.com", "https://www.moneygeek.com/insurance/business/food/cost/"),
        ("Food Business Insurance California", "FLIP", "https://www.fliprogram.com/states-page/food-liability-insurance-california"),
        ("Guide to Commercial Car Insurance in California (2026)", "Insurify", "https://insurify.com/business-insurance/commercial-auto/california/"),
        ("Food Truck Business Insurance Cost (2026 Rates)", "MoneyGeek.com", "https://www.moneygeek.com/insurance/business/food/food-truck/cost/"),
        ("Food Truck Business Insurance Costs", "Insureon", "https://www.insureon.com/food-business-insurance/food-trucks/cost"),
        ("How Much Does A Food Truck Cost in 2025? A Complete Guide", "Square", "https://squareup.com/us/en/the-bottom-line/operating-your-business/food-truck-cost"),
        ("Minimum Wage", "California Department of Industrial Relations", "https://www.dir.ca.gov/dlse/minimum_wage.htm"),
        ("2026 California Minimum Wage by City and County", "Paycor", "https://www.paycor.com/resource-center/articles/california-minimum-wage/"),
        ("Minimum Wage", "City of Los Angeles", "https://wagesla.lacity.gov/"),
        ("Minimum Wage for Workers", "Consumer and Business Affairs, LA County", "https://dcba.lacounty.gov/workers/"),
    ]

    for i, (title, source, url) in enumerate(citations, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

        # Number
        num_run = p.add_run(f"{i}. ")
        num_run.font.name = FONT_BODY
        num_run.font.size = Pt(9)
        num_run.font.color.rgb = ACCENT
        num_run.bold = True

        # Title
        title_run = p.add_run(f"\u201c{title}.\u201d ")
        title_run.font.name = FONT_BODY
        title_run.font.size = Pt(9)
        title_run.font.color.rgb = DARK_GRAY

        # Source
        source_run = p.add_run(f"{source}. ")
        source_run.font.name = FONT_BODY
        source_run.font.size = Pt(9)
        source_run.font.color.rgb = MID_GRAY
        source_run.italic = True

        # URL
        url_run = p.add_run(url)
        url_run.font.name = FONT_BODY
        url_run.font.size = Pt(8)
        url_run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

    doc.save("output/Armageddon_Treats_Works_Cited.docx")
    print("Works Cited saved: output/Armageddon_Treats_Works_Cited.docx")


if __name__ == "__main__":
    build_works_cited()
