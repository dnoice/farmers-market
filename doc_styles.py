"""
Shared document styling and helpers for all Armageddon Treats Word document modules.
Provides consistent brand identity, fonts, colors, tables, and layout utilities.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Brand Palette ─────────────────────────────────────────────
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY   = RGBColor(0x2D, 0x2D, 0x2D)
MID_GRAY    = RGBColor(0x4A, 0x4A, 0x4A)
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)
ACCENT      = RGBColor(0xC0, 0x39, 0x2B)
ACCENT_DARK = RGBColor(0x96, 0x2D, 0x22)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
OLIVE       = RGBColor(0x55, 0x6B, 0x2F)
AMBER       = RGBColor(0xD4, 0x8B, 0x0B)

TABLE_HEADER_BG = "1A1A1A"
TABLE_ALT_BG    = "F7F7F7"
TABLE_BORDER    = "CCCCCC"
ACCENT_HEX      = "C0392B"

FONT_BODY    = "Calibri"
FONT_HEADING = "Calibri"


# ── Cell-Level Helpers ────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


# ── Table Styling ─────────────────────────────────────────────
def style_header_row(row, bg=TABLE_HEADER_BG, font_color=WHITE):
    for cell in row.cells:
        set_cell_shading(cell, bg)
        set_cell_borders(cell, "999999")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = font_color if isinstance(font_color, RGBColor) else WHITE
                run.font.size = Pt(9.5)
                run.font.name = FONT_HEADING


def style_data_row(row, idx):
    bg = TABLE_ALT_BG if idx % 2 == 0 else "FFFFFF"
    for cell in row.cells:
        set_cell_shading(cell, bg)
        set_cell_borders(cell, TABLE_BORDER)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = FONT_BODY
                run.font.color.rgb = DARK_GRAY


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
        style_data_row(row, r_idx)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")
    return table


# ── Paragraph Helpers ─────────────────────────────────────────
def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_HEADING
        run.font.color.rgb = BLACK if level == 1 else DARK_GRAY
    return h


def add_body_para(doc, text, bold=False, italic=False, space_after=Pt(8)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GRAY
    run.bold = bold
    run.italic = italic
    return p


def add_bullet_point(doc, text, bold_prefix=None, indent_level=0):
    """Add a bullet-point paragraph with optional bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))

    if bold_prefix:
        run_b = p.add_run(bold_prefix + "  ")
        run_b.font.name = FONT_BODY
        run_b.font.size = Pt(10)
        run_b.font.color.rgb = DARK_GRAY
        run_b.bold = True

    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GRAY
    return p


def add_callout_box(doc, title, text, accent_color=ACCENT_HEX):
    """Add a styled callout/highlight box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FEF5F4")
    set_cell_borders(cell, accent_color, sz="8")

    p_title = cell.paragraphs[0]
    run_t = p_title.add_run(title)
    run_t.font.name = FONT_HEADING
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = ACCENT
    run_t.bold = True

    p_body = cell.add_paragraph()
    run_b = p_body.add_run(text)
    run_b.font.name = FONT_BODY
    run_b.font.size = Pt(9.5)
    run_b.font.color.rgb = MID_GRAY

    for row in table.rows:
        for c in row.cells:
            c.width = Inches(6.5)

    doc.add_paragraph("")
    return table


def add_checklist_item(doc, text, is_checked=False):
    """Add a checklist-style item with a box character."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.3)

    box = "\u2611 " if is_checked else "\u2610 "
    run_box = p.add_run(box)
    run_box.font.name = "Segoe UI Symbol"
    run_box.font.size = Pt(11)
    run_box.font.color.rgb = ACCENT if not is_checked else OLIVE

    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GRAY
    return p


# ── Document Setup ────────────────────────────────────────────
def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10)
    style.font.color.rgb = MID_GRAY

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = FONT_HEADING
        hs.font.color.rgb = BLACK if i == 1 else DARK_GRAY
        hs.paragraph_format.space_before = Pt(18 if i == 1 else 14)
        hs.paragraph_format.space_after = Pt(6)
        if i == 1:
            hs.font.size = Pt(18)
        elif i == 2:
            hs.font.size = Pt(14)
        else:
            hs.font.size = Pt(12)

    return doc


def add_header_footer(doc, title_short, footer_text=None):
    ft = footer_text or "ARMAGEDDON TREATS  |  Veteran-Owned  |  Confidential"
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(title_short)
        run.font.name = FONT_HEADING
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.small_caps = True

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(ft)
        run.font.name = FONT_HEADING
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        run.font.small_caps = True


def add_module_cover(doc, module_number, module_title, subtitle, meta_lines=None):
    """Create a branded module cover page."""
    for _ in range(4):
        doc.add_paragraph("")

    # Module number badge
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run(f"MODULE {module_number}")
    run0.font.name = FONT_HEADING
    run0.font.size = Pt(12)
    run0.font.color.rgb = ACCENT
    run0.bold = True
    run0.font.small_caps = True

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(module_title.upper())
    run.font.name = FONT_HEADING
    run.font.size = Pt(28)
    run.font.color.rgb = BLACK
    run.bold = True
    run.font.small_caps = True

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    run2 = p2.add_run(subtitle)
    run2.font.name = FONT_HEADING
    run2.font.size = Pt(14)
    run2.font.color.rgb = MID_GRAY

    # Divider
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    run3 = p3.add_run("_" * 55)
    run3.font.color.rgb = ACCENT
    run3.font.size = Pt(10)

    # Brand line
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(20)
    run4 = p4.add_run("ARMAGEDDON TREATS")
    run4.font.name = FONT_HEADING
    run4.font.size = Pt(11)
    run4.font.color.rgb = DARK_GRAY
    run4.font.small_caps = True

    # Meta info
    if meta_lines:
        for _ in range(1):
            doc.add_paragraph("")
        for line in meta_lines:
            pm = doc.add_paragraph()
            pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pm.paragraph_format.space_after = Pt(2)
            run_m = pm.add_run(line)
            run_m.font.name = FONT_BODY
            run_m.font.size = Pt(9.5)
            run_m.font.color.rgb = MID_GRAY

    doc.add_page_break()


def add_phase_banner(doc, phase_num, phase_title, date_range, color_hex=ACCENT_HEX):
    """Add a prominent phase/milestone banner."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell_left = table.rows[0].cells[0]
    cell_right = table.rows[0].cells[1]

    set_cell_shading(cell_left, color_hex)
    set_cell_shading(cell_right, color_hex)
    set_cell_borders(cell_left, color_hex, "6")
    set_cell_borders(cell_right, color_hex, "6")

    p_left = cell_left.paragraphs[0]
    run_l = p_left.add_run(f"PHASE {phase_num}: {phase_title}")
    run_l.font.name = FONT_HEADING
    run_l.font.size = Pt(12)
    run_l.font.color.rgb = WHITE
    run_l.bold = True

    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r = p_right.add_run(date_range)
    run_r.font.name = FONT_BODY
    run_r.font.size = Pt(10)
    run_r.font.color.rgb = WHITE

    cell_left.width = Inches(4.5)
    cell_right.width = Inches(2.0)

    doc.add_paragraph("")
    return table
